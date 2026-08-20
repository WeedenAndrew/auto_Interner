"""Rebuild the public fictional resume fixture; never use personal source data."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

TARGET = (
    Path(__file__).resolve().parents[2]
    / "src"
    / "auto_interner"
    / "demo_data"
    / "fictional_base_resume.docx"
)
BLUE = RGBColor(0x2E, 0x74, 0xB5)


def _add_hyperlink(paragraph: object, text: str, url: str) -> None:
    part = paragraph.part  # type: ignore[attr-defined]
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    value = OxmlElement("w:t")
    value.text = text
    run.extend((properties, value))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)  # type: ignore[attr-defined]


def _configure_styles(document: object) -> None:
    styles = document.styles  # type: ignore[attr-defined]
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading = styles["Heading 2"]
    heading.font.name = "Calibri"
    heading.font.size = Pt(13)
    heading.font.color.rgb = BLUE
    heading.font.bold = True
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.space_after = Pt(7)

    bullet = styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25


def build() -> Path:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.right_margin = Inches(0.75)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    _configure_styles(document)

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(4)
    run = name.add_run("Jordan Example")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.color.rgb = BLUE

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(6)
    contact.add_run("jordan@example.invalid | 202-555-0147 | ")
    _add_hyperlink(contact, "portfolio.example.invalid", "https://portfolio.example.invalid")

    document.add_paragraph("Experience", style="Heading 2")
    role = document.add_paragraph()
    role.add_run("Software Engineering Intern — Northstar Fictional Labs").bold = True
    role.add_run(" | May 2026–August 2026")
    document.add_paragraph(
        "Built Python validation tools that reduced fictional review time by 30% "
        "across 12,000 records.",
        style="List Bullet",
    )
    document.add_paragraph(
        "Wrote SQL checks and Git-based tests for a fictional data pipeline used by 4 teammates.",
        style="List Bullet",
    )

    document.add_paragraph("Projects", style="Heading 2")
    project = document.add_paragraph()
    project.add_run("Campus Transit Simulator").bold = True
    project.add_run(" — Python, Flask, PostgreSQL")
    document.add_paragraph(
        "Created a fictional route simulator with 25 test scenarios and a documented API.",
        style="List Bullet",
    )
    project_link = document.add_paragraph("Fictional project reference: ")
    _add_hyperlink(project_link, "example.invalid/project", "https://example.invalid/project")

    document.add_paragraph("Education", style="Heading 2")
    education = document.add_paragraph()
    education.add_run("Example State University").bold = True
    education.add_run(" — B.S. Computer Science, expected May 2028")
    document.add_paragraph(
        "Relevant coursework: Data Structures, Algorithms, Databases, and Software Engineering."
    )

    document.add_paragraph("Technical Skills", style="Heading 2")
    document.add_paragraph(
        "Languages: Python, Java, SQL, JavaScript | Tools: Git, Docker, Linux | Frameworks: Flask"
    )

    properties = document.core_properties
    properties.author = "Fictional Fixture Builder"
    properties.last_modified_by = "Fictional Fixture Builder"
    properties.title = "Fictional Resume Fixture"
    properties.subject = "Public automated test data"
    properties.comments = "Contains no real identity or employment history."
    TARGET.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(TARGET))
    return TARGET


if __name__ == "__main__":
    print(build())
