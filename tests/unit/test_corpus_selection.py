"""F-COR corpus selection, coverage, and gap-reporting tests.

The first test is the package's entire premise expressed as an assertion: a
rendered resume may contain only lines the user wrote. If it fails, selection
has started behaving like generation and the truthfulness guarantee is gone.
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest

from auto_interner.corpus import (
    Block,
    BlockKind,
    Bullet,
    CorpusError,
    build_report,
    extract_requirements,
    load_corpus,
    render_resume,
    select,
)
from auto_interner.corpus.requirements import Priority

DEMO_CORPUS = (
    Path(__file__).resolve().parents[2]
    / "src" / "auto_interner" / "demo_data" / "fictional_corpus.json"
)

BACKEND_POSTING = """
Requirements:
- Strong Python and SQL
- Experience building ETL pipelines
- Comfortable with Docker and Linux

Nice to have:
- Kubernetes and Terraform
- Kafka
"""


def _load() -> tuple[Block, ...]:
    return load_corpus(DEMO_CORPUS)


@pytest.fixture
def blocks() -> tuple[Block, ...]:
    return _load()


# ── The core invariant ───────────────────────────────────────────────────────

def test_every_rendered_line_exists_verbatim_in_the_corpus(blocks: tuple[Block, ...]) -> None:
    requirements = extract_requirements(BACKEND_POSTING)
    rendered = render_resume(select(blocks, requirements, budget=30))

    authored = {bullet.text for block in blocks for bullet in block.bullets}
    for line in rendered.splitlines():
        if line.strip().startswith("•"):
            assert line.strip()[1:].strip() in authored


# ── Requirement extraction ───────────────────────────────────────────────────

def test_nice_to_have_header_governs_every_line_beneath_it() -> None:
    by_term = {r.term: r for r in extract_requirements(BACKEND_POSTING)}
    assert by_term["python"].priority is Priority.REQUIRED
    for term in ("kubernetes", "terraform", "kafka"):
        assert by_term[term].priority is Priority.PREFERRED, term


def test_inline_cue_overrides_its_section() -> None:
    posting = "Requirements:\n- Python\n- Rust is a plus\n"
    by_term = {r.term: r for r in extract_requirements(posting)}
    assert by_term["python"].priority is Priority.REQUIRED
    assert by_term["rust"].priority is Priority.PREFERRED


# ── Selection ────────────────────────────────────────────────────────────────

def test_block_tags_do_not_suppress_later_bullets() -> None:
    """Crediting block tags per-bullet made bullet #1 appear to cover everything."""
    block = Block(
        id="b", kind=BlockKind.EXPERIENCE, title="Engineer",
        tags=frozenset({"python", "docker"}),
        bullets=(
            Bullet("first", frozenset({"python"})),
            Bullet("second", frozenset({"docker"})),
            Bullet("third", frozenset({"sql"})),
        ),
    )
    requirements = extract_requirements("Requirements:\n- python\n- docker\n- sql\n")
    assert len(select((block,), requirements, budget=20).blocks[0].bullets) == 3


def test_block_without_bullets_is_selectable() -> None:
    skills = Block(
        id="s", kind=BlockKind.SKILL, title="Python, SQL",
        tags=frozenset({"python", "sql"}), bullets=(),
    )
    requirements = extract_requirements("Requirements:\n- python\n- sql\n")
    selection = select((skills,), requirements, budget=10)
    assert len(selection.blocks) == 1
    assert "sql" in selection.covered


def test_budget_is_spent_not_abandoned_at_coverage(blocks: tuple[Block, ...]) -> None:
    """Pure max-coverage stops once requirements are met, leaving the page blank."""
    selection = select(blocks, extract_requirements(BACKEND_POSTING), budget=26)
    assert 12 <= selection.used <= 26


def test_selection_keeps_different_blocks_for_a_different_posting() -> None:
    """The premise: two postings, one corpus, different documents.

    This used the shared fixture at budget 26 against a 23-line corpus, so
    everything fitted and nothing competed — it passed on incidental ordering
    and would have gone green whatever the ranking did. The fixture has only
    three projects, so once every project must show two bullets and at most
    three may appear, it can no longer force a choice at any budget. Build the
    competition explicitly instead.
    """
    from auto_interner.corpus.selection import Shape

    backend_project = Block(
        id="api", kind=BlockKind.PROJECT, title="Service", recency=9,
        tags=frozenset({"python", "docker"}),
        bullets=(Bullet("Built a python service", frozenset({"python"})),
                 Bullet("Containerised it with docker", frozenset({"docker"}))),
    )
    mobile_project = Block(
        id="app", kind=BlockKind.PROJECT, title="Coin Flip", recency=8,
        tags=frozenset({"flutter", "dart"}),
        bullets=(Bullet("Shipped a flutter app", frozenset({"flutter"})),
                 Bullet("Wrote unit tests for it", frozenset({"unit testing"}))),
    )
    corpus = (backend_project, mobile_project)
    shape = Shape(min_experience=0, max_blocks_by_kind={"project": 1})

    backend = select(corpus, extract_requirements(BACKEND_POSTING), budget=12, shape=shape)
    mobile = select(corpus, extract_requirements("Requirements:\n- Flutter\n"),
                    budget=12, shape=shape)

    assert [b.block.id for b in backend.blocks] == ["api"]
    assert [b.block.id for b in mobile.blocks] == ["app"]


# ── Coverage and gaps ────────────────────────────────────────────────────────

def test_unsupported_requirement_is_reported_not_filled(blocks: tuple[Block, ...]) -> None:
    requirements = extract_requirements("Requirements:\n- Python\n- Kafka\n- Airflow\n")
    report = build_report(blocks, requirements, select(blocks, requirements, budget=30))
    gaps = {status.requirement.term for status in report.gaps}
    assert {"kafka", "airflow"} <= gaps
    assert "python" not in gaps


def test_score_is_bounded(blocks: tuple[Block, ...]) -> None:
    requirements = extract_requirements(BACKEND_POSTING)
    report = build_report(blocks, requirements, select(blocks, requirements, budget=26))
    assert 0.0 <= report.score() <= 1.0


# ── Corpus validation ────────────────────────────────────────────────────────

def test_duplicate_ids_rejected(tmp_path: Path) -> None:
    path = tmp_path / "corpus.json"
    path.write_text(json.dumps({"blocks": [
        {"id": "a", "kind": "project", "title": "X"},
        {"id": "a", "kind": "project", "title": "Y"},
    ]}), encoding="utf-8")
    with pytest.raises(CorpusError, match="Duplicate"):
        load_corpus(path)


def test_unknown_kind_rejected(tmp_path: Path) -> None:
    path = tmp_path / "corpus.json"
    path.write_text(json.dumps({"blocks": [{"id": "a", "kind": "nonsense", "title": "X"}]}),
                    encoding="utf-8")
    with pytest.raises(CorpusError, match="kind must be one of"):
        load_corpus(path)


def test_malformed_json_rejected(tmp_path: Path) -> None:
    path = tmp_path / "corpus.json"
    path.write_text("{not json", encoding="utf-8")
    with pytest.raises(CorpusError, match="not valid JSON"):
        load_corpus(path)


def test_missing_corpus_rejected(tmp_path: Path) -> None:
    with pytest.raises(CorpusError, match="Corpus not found"):
        load_corpus(tmp_path / "absent.json")


def test_metrics_extracted_for_tamper_detection() -> None:
    assert Bullet("cut latency 40% across 3 services").metrics == ("40%", "3")


def test_budget_is_filled_even_when_nothing_else_matches_the_posting() -> None:
    """Affinity fill stopping early left two thirds of a resume blank.

    Education carries no tag any posting asks for, so it scored zero and was
    dropped — from an internship resume, where omitting it is disqualifying.
    """
    education = Block(
        id="edu", kind=BlockKind.EDUCATION, title="BSc Computer Science",
        recency=9, bullets=(Bullet("Expected May 2028"),),
    )
    relevant = Block(
        id="job", kind=BlockKind.EXPERIENCE, title="Engineer",
        recency=10, tags=frozenset({"python"}),
        bullets=(Bullet("Shipped a thing", frozenset({"python"})),),
    )
    requirements = extract_requirements("Requirements:\n- Python\n")
    selection = select((education, relevant), requirements, budget=20)

    assert {s.block.id for s in selection.blocks} == {"edu", "job"}


def test_education_is_included_in_full_regardless_of_the_posting() -> None:
    """No posting lists a degree as a requirement, so coverage always drops it."""
    from auto_interner.corpus.selection import Shape

    education = Block(
        id="edu", kind=BlockKind.EDUCATION, title="BSc Computer Science",
        bullets=(Bullet("Expected May 2028"), Bullet("Coursework: Data Structures")),
    )
    job = Block(
        id="j", kind=BlockKind.EXPERIENCE, title="Engineer", recency=9,
        tags=frozenset({"python"}), bullets=(Bullet("Shipped", frozenset({"python"})),),
    )
    selection = select((education, job), extract_requirements("Requirements:\n- Python\n"),
                       budget=30, shape=Shape(min_experience=0))
    edu = next(s for s in selection.blocks if s.block.id == "edu")
    assert len(edu.bullets) == 2, "education must appear in full, not trimmed"


def test_at_least_two_experience_entries_are_kept() -> None:
    """One job reads as no history, whatever the posting asked for."""
    from auto_interner.corpus.selection import Shape

    jobs = tuple(
        Block(id=f"j{i}", kind=BlockKind.EXPERIENCE, title=f"Role {i}", recency=10 - i,
              bullets=(Bullet(f"did thing {i}"),))
        for i in range(3)
    )
    selection = select(jobs, extract_requirements("Requirements:\n- Kubernetes\n"),
                       budget=30, shape=Shape(min_experience=2))
    kept = [s for s in selection.blocks if s.block.kind is BlockKind.EXPERIENCE]
    assert len(kept) >= 2
    assert [s.block.id for s in kept][:2] == ["j0", "j1"], "most recent first"


def test_underfill_is_reported_rather_than_padded() -> None:
    """A thin resume is a corpus problem. Say so; do not widen the margins."""
    from auto_interner.corpus.selection import Shape

    thin = Block(id="a", kind=BlockKind.PROJECT, title="One thing",
                 tags=frozenset({"python"}),
                 bullets=(Bullet("did a thing", frozenset({"python"})),))
    selection = select((thin,), extract_requirements("Requirements:\n- Python\n"),
                       budget=40, shape=Shape(min_experience=0))

    assert selection.underfilled
    assert "corpus is exhausted, not the page" in selection.fill_report()


def test_a_full_page_is_not_reported_as_short() -> None:
    from auto_interner.corpus.selection import Shape

    blocks = tuple(
        Block(id=f"b{i}", kind=BlockKind.PROJECT, title=f"Project {i}",
              recency=i, tags=frozenset({"python"}),
              bullets=(Bullet(f"built thing {i}", frozenset({"python"})),))
        for i in range(6)
    )
    selection = select(blocks, extract_requirements("Requirements:\n- Python\n"),
                       budget=12, shape=Shape(min_experience=0, max_blocks_by_kind={}))
    assert not selection.underfilled
    assert "full" in selection.fill_report()


# ── template assembly ────────────────────────────────────────────────────────

def _blocks_from_resume(path: Path) -> tuple[Block, ...]:
    """Build a corpus out of a resume, which is how the real flow works.

    `fictional_corpus.json` and `fictional_base_resume.docx` are unrelated
    fixtures - zero bullet overlap - so pairing them tests nothing. Template
    assembly only makes sense when the corpus came from the document.
    """
    from auto_interner.documents.template_reader import read_resume

    doc = read_resume(path)
    blocks: list[Block] = []
    for section in doc.sections:
        kind = {"experience": BlockKind.EXPERIENCE, "projects": BlockKind.PROJECT,
                "education": BlockKind.EDUCATION}.get(
                    section.name.casefold(), BlockKind.SKILL)
        bullets = tuple(Bullet(p.source_text, frozenset({"python"}))
                        for p in section.paragraphs if len(p.source_text) > 40)
        if bullets:
            blocks.append(Block(id=section.name.casefold().replace(" ", "-"),
                                kind=kind, title=section.name,
                                tags=frozenset({"python"}), bullets=bullets))
    return tuple(blocks)


def test_template_assembly_keeps_only_selected_content(tmp_path: Path) -> None:
    """Output must be the user's own document with paragraphs removed.

    Building a fresh DOCX produces something correct that looks nothing like
    the resume they have been sending, which makes it a new document to
    proofread rather than a tailored one.
    """
    from auto_interner.corpus.assemble import assemble_from_template
    from auto_interner.documents.template_reader import read_resume

    base = (
        Path(__file__).resolve().parents[2]
        / "src" / "auto_interner" / "demo_data" / "fictional_base_resume.docx"
    )
    source = read_resume(base)
    blocks = _blocks_from_resume(base)
    selection = select(blocks, extract_requirements("Requirements:\n- Python\n"), budget=40)

    result = assemble_from_template(source, selection, tmp_path / "out.docx")
    assert result.destination.is_file()
    assert result.kept > 0
    assert result.sections_kept, "at least one section must survive"


def test_template_output_invents_nothing(tmp_path: Path) -> None:
    """Every line in the output must already exist in the source document."""
    from docx import Document

    from auto_interner.corpus.assemble import assemble_from_template
    from auto_interner.documents.template_reader import read_resume

    base = (
        Path(__file__).resolve().parents[2]
        / "src" / "auto_interner" / "demo_data" / "fictional_base_resume.docx"
    )
    source = read_resume(base)
    selection = select(_blocks_from_resume(base),
                       extract_requirements("Requirements:\n- Python\n"), budget=40)
    out = tmp_path / "out.docx"
    assemble_from_template(source, selection, out)

    original = {p.text.strip() for p in Document(str(base)).paragraphs}
    for paragraph in Document(str(out)).paragraphs:
        text = paragraph.text.strip()
        if text:
            assert text in original


def test_template_assembly_refuses_to_overwrite(tmp_path: Path) -> None:
    from auto_interner.corpus.assemble import TemplateAssemblyError, assemble_from_template
    from auto_interner.documents.template_reader import read_resume

    base = (
        Path(__file__).resolve().parents[2]
        / "src" / "auto_interner" / "demo_data" / "fictional_base_resume.docx"
    )
    dest = tmp_path / "out.docx"
    dest.write_text("already here", encoding="utf-8")
    with pytest.raises(TemplateAssemblyError, match="refusing to overwrite"):
        assemble_from_template(
            read_resume(base),
            select(_blocks_from_resume(base), extract_requirements("Requirements:\n- Python\n")),
            dest,
        )


# --- extraction boundaries -------------------------------------------------
#
# All three of these were found by running the extractor against a live
# ByteDance infrastructure posting rather than a fixture. Each one failed
# silently: the requirement simply never appeared, so it was absent from the
# gap list too and the resume scored as a fuller match than it was.


def test_a_term_ending_a_sentence_is_still_extracted() -> None:
    """The period is sentence punctuation, not part of the term."""
    assert [r.term for r in extract_requirements("Experience with Docker.")] == ["docker"]


def test_a_dotted_term_is_not_split_by_the_sentence_rule() -> None:
    """Relaxing the period must not make 'node' match inside 'node.js'."""
    terms = {r.term for r in extract_requirements("We use Node.js here.")}
    assert "node.js" in terms


def test_every_language_in_a_closing_list_is_extracted() -> None:
    """'Java, Go, C++, or Python.' must not lose the last language."""
    posting = "Strong programming skills in one or more of Java, Go, C++, or Python."
    assert {r.term for r in extract_requirements(posting)} == {"java", "go", "c++", "python"}


def test_concept_requirements_are_recognised_without_a_named_tool() -> None:
    """Infrastructure postings name capabilities, not products."""
    posting = (
        "Requirements:\n"
        "- Build high availability and disaster recovery capabilities.\n"
        "- Own traffic routing and improve system reliability at scale.\n"
    )
    terms = {r.term for r in extract_requirements(posting)}
    assert {"high availability", "disaster recovery", "traffic routing", "reliability"} <= terms


def test_a_colonless_qualifications_header_governs_its_section() -> None:
    """'Preferred Qualifications' with no colon is the common house style."""
    posting = (
        "Minimum Qualifications\n"
        "- Strong programming skills in Python.\n"
        "Preferred Qualifications\n"
        "- Good understanding of machine learning.\n"
    )
    priorities = {r.term: r.priority for r in extract_requirements(posting)}
    assert priorities["python"] is Priority.REQUIRED
    assert priorities["machine learning"] is Priority.PREFERRED


def test_a_sentence_beginning_with_a_cue_is_not_eaten_as_a_header() -> None:
    """A colonless header must look like a heading, not a requirement line."""
    terms = {r.term for r in extract_requirements("Preferred experience with Kafka\n")}
    assert "kafka" in terms


# --- evidence subsumption --------------------------------------------------


def test_a_specific_tool_evidences_the_general_capability() -> None:
    """A posting says 'databases'; a resume says 'MongoDB'."""
    assert "databases" in Bullet("used mongo", frozenset({"mongodb"})).tags
    assert "containers" in Block(id="b", kind=BlockKind.PROJECT, title="t",
                                 tags=frozenset({"docker"})).tags


def test_subsumption_keeps_the_original_tag() -> None:
    """Broadening must add, never replace: 'mongodb' is still matchable."""
    assert {"mongodb", "databases"} <= Bullet("x", frozenset({"mongodb"})).tags


def test_subsumption_does_not_invent_adjacent_skills() -> None:
    """Adjacent is not implied. Overstating coverage hides real gaps."""
    docker = Bullet("x", frozenset({"docker"})).tags
    assert "distributed systems" not in docker
    assert "operating systems" not in Bullet("x", frozenset({"linux"})).tags


def test_a_subsumed_requirement_is_not_reported_as_a_gap() -> None:
    """The whole point: 'databases' must not appear as unsupported."""
    blocks = (
        Block(id="e", kind=BlockKind.EDUCATION, title="Uni", tags=frozenset({"python"}),
              bullets=(Bullet("studied things", frozenset({"python"})),)),
        Block(id="j", kind=BlockKind.EXPERIENCE, title="Job", tags=frozenset({"mongodb"}),
              bullets=(Bullet("stored records in mongodb", frozenset({"mongodb"})),)),
    )
    from auto_interner.corpus.selection import Shape

    requirements = extract_requirements("Requirements:\n- Experience with databases.\n")
    selection = select(blocks, requirements, budget=12, shape=Shape(min_experience=0))
    report = build_report(blocks, requirements, selection)
    assert [g.requirement.term for g in report.gaps] == []


# --- the title-line link rule ----------------------------------------------
#
# A repository link belongs at the right margin of its project's title line,
# where experience and education already put location and dates. The base
# resume got there with 84 literal spaces, which is not a position: it held at
# exactly one font size and margin, and in practice overflowed and wrapped the
# link onto a line of its own -- spending a line of a one-page budget on a URL.

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _paragraph_with_padded_link(padding: str = " " * 40):
    """A project title, whitespace padding, then a hyperlink. The real shape."""
    from docx import Document as _Document

    document = _Document()
    paragraph = document.add_paragraph("Auto Interner | Python, Docker")
    paragraph.add_run(padding)
    link = paragraph._p.makeelement(f"{_W}hyperlink", {})
    run = link.makeelement(f"{_W}r", {})
    text = run.makeelement(f"{_W}t", {})
    text.text = "auto_Interner"
    run.append(text)
    link.append(run)
    paragraph._p.append(link)
    return document, paragraph


def _tab_count(paragraph) -> int:
    return len(paragraph._p.findall(f".//{_W}tab"))


def test_a_padded_project_link_moves_onto_the_title_line() -> None:
    from auto_interner.corpus.assemble import inline_trailing_link

    _, paragraph = _paragraph_with_padded_link()
    assert _tab_count(paragraph) == 0
    assert inline_trailing_link(paragraph, 10080) is True
    # One tab in a run, one right-aligned stop in the paragraph properties.
    assert _tab_count(paragraph) == 2
    stop = paragraph._p.find(f"{_W}pPr/{_W}tabs/{_W}tab")
    assert stop.get(f"{_W}val") == "right"
    assert stop.get(f"{_W}pos") == "10080"


def test_link_inlining_changes_no_words() -> None:
    """Whitespace may move. Nothing a reader would call content may."""
    from auto_interner.corpus.assemble import inline_trailing_link

    _, paragraph = _paragraph_with_padded_link()
    before = "".join(paragraph.text.split())
    inline_trailing_link(paragraph, 10080)
    assert "".join(paragraph.text.split()) == before
    assert "auto_Interner" in paragraph.text


def test_the_rule_is_idempotent() -> None:
    """Assembling an already-corrected resume must not stack tab stops."""
    from auto_interner.corpus.assemble import inline_trailing_link

    _, paragraph = _paragraph_with_padded_link()
    inline_trailing_link(paragraph, 10080)
    after_once = _tab_count(paragraph)
    inline_trailing_link(paragraph, 10080)
    assert _tab_count(paragraph) == after_once


def test_a_multi_link_line_is_left_alone() -> None:
    """Contact links are punctuation-separated prose, not a padded title."""
    from auto_interner.corpus.assemble import inline_trailing_link

    document, paragraph = _paragraph_with_padded_link()
    second = paragraph._p.makeelement(f"{_W}hyperlink", {})
    paragraph._p.append(second)
    assert inline_trailing_link(paragraph, 10080) is False


def test_a_link_after_ordinary_prose_is_not_pushed_right() -> None:
    """Only deliberate whitespace padding signals intent to right-align."""
    from auto_interner.corpus.assemble import inline_trailing_link

    _, paragraph = _paragraph_with_padded_link(padding=" ")
    assert inline_trailing_link(paragraph, 10080) is False


def test_assembly_applies_the_rule_to_the_finished_document(tmp_path: Path) -> None:
    """The rule is part of assembly, not something a caller must remember."""
    from docx import Document as _Document

    from auto_interner.corpus.assemble import inline_trailing_links

    document, _ = _paragraph_with_padded_link()
    source = tmp_path / "base.docx"
    document.save(str(source))
    reopened = _Document(str(source))
    assert inline_trailing_links(reopened) == 1


def test_a_vestigial_empty_run_does_not_block_the_rule() -> None:
    """Word leaves text-less, tab-less runs everywhere; they mean nothing."""
    from auto_interner.corpus.assemble import inline_trailing_link

    _, paragraph = _paragraph_with_padded_link()
    link = paragraph._p.findall(f"{_W}hyperlink")[0]
    empty = paragraph._p.makeelement(f"{_W}r", {})
    link.addprevious(empty)
    assert inline_trailing_link(paragraph, 10080) is True


def test_a_tab_without_a_right_stop_is_still_broken() -> None:
    """The base resume had a tab, 84 spaces, and no stop. That is not correct.

    Reading the tab as "already positioned" made the rule skip the one
    paragraph it was written for, which is how this shipped wrong once.
    """
    from auto_interner.corpus.assemble import inline_trailing_link

    _, paragraph = _paragraph_with_padded_link()
    link = paragraph._p.findall(f"{_W}hyperlink")[0]
    run = paragraph._p.makeelement(f"{_W}r", {})
    run.append(run.makeelement(f"{_W}tab", {}))
    link.addprevious(run)
    assert inline_trailing_link(paragraph, 10080) is True
    assert paragraph._p.find(f"{_W}pPr/{_W}tabs/{_W}tab").get(f"{_W}val") == "right"


# --- ranking a feed --------------------------------------------------------
#
# Found by scoring one corpus against five live postings. Ranking by score()
# put the weakest match first, because score's denominator is how many skills
# the posting bothered to name -- a fact about the writing, not the candidate.

_VAGUE = "Requirements:\n- Proficient in Python or C++.\n"
_SPECIFIC = (
    "Requirements:\n"
    "- Strong programming in Java, Go, C++, or Python.\n"
    "- Foundation in data structures, algorithms, networking, and databases.\n"
    "- Build high availability, disaster recovery and traffic routing.\n"
)


def _report_for(posting: str):
    from auto_interner.corpus.selection import Shape

    blocks = (
        Block(id="edu", kind=BlockKind.EDUCATION, title="BSc",
              tags=frozenset({"data structures"}),
              bullets=(Bullet("Coursework: data structures", frozenset({"data structures"})),)),
        Block(id="job", kind=BlockKind.EXPERIENCE, title="Engineer",
              tags=frozenset({"python", "java", "mongodb"}),
              bullets=(Bullet("shipped python and java on mongodb",
                              frozenset({"python", "java", "mongodb"})),)),
    )
    requirements = extract_requirements(posting)
    selection = select(blocks, requirements, budget=20, shape=Shape(min_experience=0))
    return build_report(blocks, requirements, selection)


def test_score_ranks_a_vague_posting_above_a_better_match() -> None:
    """The defect, asserted. score() is not a ranking signal."""
    assert _report_for(_VAGUE).score() > _report_for(_SPECIFIC).score()


def test_evidence_ranks_the_better_match_first() -> None:
    """Absolute weight covered orders the two the way a human would."""
    assert _report_for(_VAGUE).evidence() < _report_for(_SPECIFIC).evidence()


def test_a_thin_posting_is_flagged_as_not_comparable() -> None:
    """Two extracted terms cannot support a percentage worth comparing."""
    assert _report_for(_VAGUE).is_comparable() is False
    assert _report_for(_SPECIFIC).is_comparable() is True


def test_evidence_counts_weight_not_terms() -> None:
    """A required term surfaced is worth more than a preferred one."""
    report = _report_for("Requirements:\n- Python\n\nNice to have:\n- Java\n")
    assert report.evidence() == 4  # required python (3) + preferred java (1)


# --- alternation groups ----------------------------------------------------
#
# "One or more languages such as Java, Go, C++, or Python" is one requirement
# with four acceptable answers. Counting it as four told a Python programmer
# they were missing Go and C++ by the same posting that said Python was fine.


def test_an_or_list_becomes_one_requirement() -> None:
    posting = "Requirements:\n- Strong skills in one or more of Java, Go, C++, or Python.\n"
    groups = {r.group for r in extract_requirements(posting)}
    assert len(groups) == 1
    assert groups != {""}


def test_an_and_list_stays_separate() -> None:
    """Conjunction is not alternation. All of this is genuinely wanted."""
    posting = "Requirements:\n- Foundation in algorithms, networking, and databases.\n"
    assert {r.group for r in extract_requirements(posting)} == {""}


def test_covering_one_alternate_satisfies_the_whole_group() -> None:
    from auto_interner.corpus.selection import Shape

    blocks = (
        Block(id="job", kind=BlockKind.EXPERIENCE, title="Engineer",
              tags=frozenset({"python"}),
              bullets=(Bullet("shipped python", frozenset({"python"})),)),
    )
    posting = "Requirements:\n- Strong skills in one or more of Java, Go, C++, or Python.\n"
    requirements = extract_requirements(posting)
    selection = select(blocks, requirements, budget=20, shape=Shape(min_experience=0))
    report = build_report(blocks, requirements, selection)

    assert report.unsatisfied_groups() == ()
    assert report.score() == 1.0


def test_an_alternation_counts_once_toward_the_score() -> None:
    """Four alternates plus one separate requirement is two, not five."""
    from auto_interner.corpus.selection import Shape

    blocks = (
        Block(id="job", kind=BlockKind.EXPERIENCE, title="Engineer",
              tags=frozenset({"python"}),
              bullets=(Bullet("shipped python", frozenset({"python"})),)),
    )
    posting = (
        "Requirements:\n"
        "- Strong skills in one or more of Java, Go, C++, or Python.\n"
        "- Experience with Kubernetes.\n"
    )
    requirements = extract_requirements(posting)
    selection = select(blocks, requirements, budget=20, shape=Shape(min_experience=0))
    report = build_report(blocks, requirements, selection)

    assert report.score() == pytest.approx(0.5)
    assert report.unsatisfied_groups() == ("kubernetes",)


def test_a_wrapped_requirement_line_is_read_as_one_sentence() -> None:
    """Postings wrap. Splitting on the newline hid the alternation cue.

    Roblox lists nine languages after "one or more programming languages,
    including" -- with the cue on the previous line. Read line by line, all
    nine became separate requirements and eight became gaps.
    """
    posting = (
        "Requirements:\n"
        "- Proficient in one or more programming languages, including Go,\n"
        "  Node.js, Ruby, Python, C++, and Java.\n"
    )
    requirements = extract_requirements(posting)
    assert {r.group for r in requirements} != {""}
    assert len({r.group for r in requirements}) == 1


def test_separate_bullets_do_not_merge() -> None:
    """A new bullet marker ends the previous requirement."""
    posting = "Requirements:\n- Experience with Kafka or Redis.\n- Experience with Docker.\n"
    by_term = {r.term: r for r in extract_requirements(posting)}
    assert by_term["docker"].group == ""
    assert by_term["kafka"].group == by_term["redis"].group != ""


# --- phrase tagging --------------------------------------------------------
#
# A posting asks for "code review"; a resume says "flag coding discrepancies
# and issues". Same work, different words. Matching literal terms alone called
# it a gap. This layer widens what an existing sentence counts as -- and only
# that. It never touches the sentence.


def test_a_described_activity_earns_the_named_capability() -> None:
    from auto_interner.corpus.tagging import tag_text

    tags, _ = tag_text("flag coding discrepancies and issues for tech companies")
    assert "code review" in tags


def test_a_deployment_with_recovery_earns_infrastructure_and_reliability() -> None:
    from auto_interner.corpus.tagging import tag_text

    tags, _ = tag_text(
        "Prepared a Raspberry Pi 4B deployment with automated restarts, "
        "monitoring, recovery procedures and a documented validation plan."
    )
    assert {"infrastructure", "reliability", "monitoring", "automation"} <= tags


def test_it_stops_short_of_the_claims_that_would_be_a_stretch() -> None:
    """A Pi with restarts is not multi-region failover, and must not claim to be.

    Over-tagging hides real gaps, which is worse than under-tagging: the user
    stops seeing what they actually need to go and learn.
    """
    from auto_interner.corpus.tagging import tag_text

    tags, _ = tag_text(
        "Prepared a Raspberry Pi 4B deployment with automated restarts, "
        "monitoring, recovery procedures and a documented validation plan."
    )
    assert "disaster recovery" not in tags
    assert "high availability" not in tags
    assert "distributed systems" not in tags


def test_every_tag_names_the_words_that_earned_it() -> None:
    """A tag the user cannot audit is a tag they cannot disagree with."""
    from auto_interner.corpus.tagging import tag_text

    text = "flag coding discrepancies and issues"
    _, hits = tag_text(text)
    for hit in hits:
        if hit.rule == "implied":
            continue
        assert hit.matched, hit
        assert hit.matched.casefold() in text.casefold()


def test_phrase_tagging_can_be_turned_off() -> None:
    """The literal-term path stays available and unchanged."""
    from auto_interner.corpus.tagging import tag_text

    tags, _ = tag_text("flag coding discrepancies and issues", phrases=False)
    assert "code review" not in tags


def test_tagging_never_alters_the_text() -> None:
    """The premise of the package, restated where it is easiest to break."""
    from auto_interner.corpus.tagging import tag_text

    text = "Utilized Docker to run AI models locally to create and run unit tests."
    before = text
    tag_text(text)
    assert text == before
    assert Bullet(text, tag_text(text)[0]).text == before


def test_a_skills_line_can_be_filled_in_when_the_page_has_room() -> None:
    """A block with no bullets is a header and nothing else.

    The fill fallback required at least one bullet, so a skills line could
    never enter it: a resume with three spare lines went out without its
    skills section, while reporting the corpus exhausted rather than the page.
    """
    from auto_interner.corpus.selection import Shape

    job = Block(
        id="job", kind=BlockKind.EXPERIENCE, title="Engineer", recency=9,
        tags=frozenset({"python"}),
        bullets=(Bullet("shipped python", frozenset({"python"})),),
    )
    skills = Block(
        id="skills", kind=BlockKind.SKILL, title="Docker, SQL, Git",
        recency=5, tags=frozenset({"docker", "sql", "git"}), bullets=(),
    )
    requirements = extract_requirements("Requirements:\n- Python\n- Docker\n")
    selection = select((job, skills), requirements, budget=20, shape=Shape(min_experience=0))
    assert "skills" in {s.block.id for s in selection.blocks}


def test_git_evidences_version_control() -> None:
    """Nearly every posting asks for it, and it was not in the taxonomy at all."""
    assert "version control" in Bullet("x", frozenset({"git"})).tags
    extracted = extract_requirements("Requirements:\n- Comfortable with Git.\n")
    assert "git" in {r.term for r in extracted}


def test_a_bulletless_project_is_not_filled_back_in() -> None:
    """A project title with nothing under it is an unfinished entry.

    Allowing every bulletless block into the fill fallback -- the fix that let
    skills lines in -- also put bare project titles back on the page, which is
    the exact defect the drop pass exists to remove.
    """
    from auto_interner.corpus.selection import Shape

    job = Block(id="job", kind=BlockKind.EXPERIENCE, title="Engineer", recency=9,
                tags=frozenset({"python"}),
                bullets=(Bullet("shipped python", frozenset({"python"})),))
    empty = Block(id="stub", kind=BlockKind.PROJECT, title="Half-Finished Thing",
                  recency=8, tags=frozenset({"python"}), bullets=())
    requirements = extract_requirements("Requirements:\n- Python\n")
    selection = select((job, empty), requirements, budget=30, shape=Shape(min_experience=0))
    assert "stub" not in {s.block.id for s in selection.blocks}


def test_git_answers_a_posting_that_names_github() -> None:
    """The transferable skill is Git. The host is not the requirement."""
    assert "github" in Bullet("x", frozenset({"git"})).tags


def test_a_tie_is_broken_by_relevance_not_by_length() -> None:
    """Two bullets covering nothing new: the on-topic one should win.

    The old rule preferred the shorter bullet outright. Against a live
    reliability posting that happened to keep the right one, which is worse
    than failing: it looked like judgment and was arithmetic on string length.
    """
    from auto_interner.corpus.selection import Shape

    long_and_relevant = Bullet(
        "Prepared a deployment with automated restarts and monitoring, plus a "
        "documented validation plan covering every recovery path in the system",
        frozenset({"monitoring", "reliability", "infrastructure"}),
    )
    short_and_not = Bullet("Wrote some docs", frozenset())
    block = Block(
        id="p", kind=BlockKind.PROJECT, title="Thing", tags=frozenset({"python"}),
        bullets=(long_and_relevant, short_and_not),
    )
    requirements = extract_requirements(
        "Requirements:\n- Python\n- Monitoring, reliability and infrastructure.\n"
    )
    selection = select((block,), requirements, budget=6, shape=Shape(min_experience=0))
    kept = [b.bullet.text for b in selection.blocks[0].bullets]
    assert kept[0] is long_and_relevant.text or kept[0] == long_and_relevant.text


def test_author_order_breaks_a_genuine_tie() -> None:
    """When nothing distinguishes two bullets, respect how the user ordered them."""
    from auto_interner.corpus.selection import Shape

    first = Bullet("Did the thing the user put first", frozenset())
    second = Bullet("Second", frozenset())
    block = Block(
        id="p", kind=BlockKind.PROJECT, title="Thing", tags=frozenset({"python"}),
        bullets=(first, second),
    )
    requirements = extract_requirements("Requirements:\n- Python\n")
    selection = select((block,), requirements, budget=4, shape=Shape(min_experience=0))
    assert selection.blocks[0].bullets[0].bullet.text == first.text


def test_a_mobile_posting_can_see_a_flutter_project() -> None:
    """"Mobile development" was not a term, so mobile work was invisible.

    A posting asking for iOS or Android named no framework the corpus held,
    and the one mobile project lost the page to a backend one.
    """
    from auto_interner.corpus.selection import Shape

    mobile = Block(
        id="app", kind=BlockKind.PROJECT, title="Coin Flip", recency=8,
        tags=frozenset({"flutter", "dart"}),
        bullets=(Bullet("Shipped an animated game", frozenset({"flutter"})),),
    )
    backend = Block(
        id="api", kind=BlockKind.PROJECT, title="Service", recency=9,
        tags=frozenset({"python"}),
        bullets=(Bullet("Built a service", frozenset({"python"})),),
    )
    posting = (
        "Minimum Qualifications\n"
        "- Proficiency in at least one mobile development language such as\n"
        "  Java, Kotlin, Objective-C, or Swift.\n"
        "- Core concepts of iOS and/or Android platforms.\n"
    )
    requirements = extract_requirements(posting)
    assert "mobile development" in {r.term for r in requirements}

    selection = select((mobile, backend), requirements, budget=4,
                       shape=Shape(min_experience=0))
    assert "app" in {s.block.id for s in selection.blocks}


def test_a_term_seen_earlier_still_joins_a_later_alternation() -> None:
    """A group must hold every term its own label names.

    Bailing out on the first sighting of a term meant it could never join a
    later alternation, so the group was reported unsatisfied while one of the
    terms printed in its name was plainly covered.
    """
    posting = (
        "Requirements:\n"
        "- Participate in implementation and debugging of features.\n"
        "- Understanding of concepts such as app lifecycle and debugging.\n"
    )
    by_term = {r.term: r for r in extract_requirements(posting)}
    assert by_term["debugging"].group
    assert by_term["debugging"].group == by_term["app lifecycle"].group


# --- shape: two bullets per project ----------------------------------------


def test_every_project_shows_two_bullets_when_it_has_them() -> None:
    """Coverage alone gives the projects section a ragged shape.

    One project earns four bullets while the next renders as a bare title,
    because its second bullet covered nothing new. That reads as an abandoned
    entry rather than a deliberate one.
    """
    from auto_interner.corpus.selection import Shape

    projects = tuple(
        Block(id=f"p{i}", kind=BlockKind.PROJECT, title=f"Project {i}", recency=9 - i,
              tags=frozenset({"python"}),
              bullets=(Bullet(f"first thing {i}", frozenset({"python"})),
                       Bullet(f"second thing {i}", frozenset()),
                       Bullet(f"third thing {i}", frozenset())))
        for i in range(3)
    )
    requirements = extract_requirements("Requirements:\n- Python\n")
    selection = select(projects, requirements, budget=40, shape=Shape(min_experience=0))
    shown = {s.block.id: len(s.bullets) for s in selection.blocks}
    assert shown and all(count == 2 for count in shown.values()), shown


def test_a_project_with_one_bullet_still_shows_one() -> None:
    """A floor, not padding. There is nothing to pad with."""
    from auto_interner.corpus.selection import Shape

    thin = Block(id="p", kind=BlockKind.PROJECT, title="Thin", tags=frozenset({"python"}),
                 bullets=(Bullet("only thing", frozenset({"python"})),))
    selection = select((thin,), extract_requirements("Requirements:\n- Python\n"),
                       budget=20, shape=Shape(min_experience=0))
    assert len(selection.blocks[0].bullets) == 1


def test_the_rule_is_configurable_and_off_by_default_for_other_kinds() -> None:
    from auto_interner.corpus.selection import Shape

    job = Block(id="j", kind=BlockKind.EXPERIENCE, title="Job", tags=frozenset({"python"}),
                bullets=(Bullet("a", frozenset({"python"})), Bullet("b", frozenset()),
                         Bullet("c", frozenset())))
    selection = select((job,), extract_requirements("Requirements:\n- Python\n"),
                       budget=20, shape=Shape(min_experience=0, bullets_by_kind={}))
    assert len(selection.blocks[0].bullets) == 1


# --- one-page fitting ------------------------------------------------------


def _doc_with(lines: int):
    from docx import Document as _Document

    document = _Document()
    for i in range(lines):
        document.add_paragraph(f"Line {i} " + "x" * 60)
    return document


def test_a_document_that_already_fits_is_left_alone() -> None:
    from auto_interner.corpus.formatting import current_typography, fit_to_one_page

    document = _doc_with(10)
    before = current_typography(document)
    result = fit_to_one_page(document)
    assert result.fitted
    assert result.final == before


def test_an_overflowing_document_is_tightened_until_it_fits() -> None:
    """55 lines is over a default page and inside what the floors allow."""
    from auto_interner.corpus.formatting import fit_to_one_page

    document = _doc_with(55)
    result = fit_to_one_page(document)
    assert result.fitted, result.describe()
    assert result.steps
    assert result.lines <= result.capacity


def test_it_stops_at_the_floors_and_says_so() -> None:
    """An unreadable resume that fits is not a success."""
    from auto_interner.corpus.formatting import Floors, fit_to_one_page

    document = _doc_with(400)
    result = fit_to_one_page(document, floors=Floors())
    assert not result.fitted
    assert result.final.font_pt >= Floors().font_pt
    assert result.final.margin_in >= Floors().margin_in
    assert "STILL OVER" in result.describe()


def test_type_size_is_the_last_thing_touched() -> None:
    """A reader notices 9.5pt type long before a 0.6in margin."""
    from auto_interner.corpus.formatting import fit_to_one_page

    result = fit_to_one_page(_doc_with(58))
    joined = " | ".join(result.steps)
    if "font ->" in joined and "margins ->" in joined:
        assert joined.index("margins ->") < joined.index("font ->")


def test_fitting_never_changes_a_word() -> None:
    """Typography only. The premise again, where it is easiest to break."""
    from auto_interner.corpus.formatting import fit_to_one_page

    document = _doc_with(55)
    before = [p.text for p in document.paragraphs if p.text.strip()]
    fit_to_one_page(document)
    assert [p.text for p in document.paragraphs if p.text.strip()] == before


def test_only_two_projects_reach_the_page() -> None:
    """Five projects reads as a list of hobbies, not a claim about strengths.

    Once the budget came from real page capacity, everything fitted — and with
    room for everything, nothing competes and two postings produce the same
    document. The cap is what keeps selection selecting.

    Three fit the line estimate and still spilled in Word, leaving one bullet
    alone on page two.
    """
    from auto_interner.corpus.selection import Shape

    projects = tuple(
        Block(id=f"p{i}", kind=BlockKind.PROJECT, title=f"Project {i}", recency=9 - i,
              tags=frozenset({"python"}),
              bullets=(Bullet(f"one {i}", frozenset({"python"})),
                       Bullet(f"two {i}", frozenset())))
        for i in range(6)
    )
    selection = select(projects, extract_requirements("Requirements:\n- Python\n"),
                       budget=60, shape=Shape(min_experience=0))
    kept = [s for s in selection.blocks if s.block.kind is BlockKind.PROJECT]
    assert len(kept) == 2
    assert all(len(s.bullets) == 2 for s in kept)


def test_the_fill_pass_does_not_deepen_one_project_past_the_others() -> None:
    """Uniformity is the point; leftover budget must not break it."""
    from auto_interner.corpus.selection import Shape

    rich = Block(id="rich", kind=BlockKind.PROJECT, title="Rich", recency=9,
                 tags=frozenset({"python"}),
                 bullets=tuple(Bullet(f"thing {i}", frozenset({"python"})) for i in range(5)))
    lean = Block(id="lean", kind=BlockKind.PROJECT, title="Lean", recency=8,
                 tags=frozenset({"python"}),
                 bullets=(Bullet("a", frozenset({"python"})), Bullet("b", frozenset())))
    selection = select((rich, lean), extract_requirements("Requirements:\n- Python\n"),
                       budget=60, shape=Shape(min_experience=0))
    assert {len(s.bullets) for s in selection.blocks} == {2}


def test_capacity_is_biased_toward_underfilling() -> None:
    """The two failures are not symmetric.

    Under-filling costs white space. Overflowing costs a second sheet carrying
    one orphaned line. The model agreed with LibreOffice exactly and still
    spilled in Word, so the estimate is deliberately under the arithmetic
    maximum.
    """
    from docx import Document as _Document

    from auto_interner.corpus.formatting import capacity, current_typography

    document = _Document()
    typography = current_typography(document)
    assert capacity(document, typography) < capacity(document, typography, safety=1.0)
