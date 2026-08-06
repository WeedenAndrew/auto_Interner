"""PII boundary and stable resume-structure extraction tests."""

from __future__ import annotations

import json
import shutil
from pathlib import Path

import pytest
from docx import Document

from auto_interner.documents.template_reader import (
    ResumeStructureError,
    contains_pii,
    read_resume,
    redact_pii,
)

pytestmark = [pytest.mark.unit, pytest.mark.privacy]
FIXTURE = (
    Path(__file__).resolve().parents[2]
    / "src"
    / "auto_interner"
    / "demo_data"
    / "fictional_base_resume.docx"
)


def _copy_fixture(tmp_path: Path) -> Path:
    destination = tmp_path / "base.docx"
    shutil.copyfile(FIXTURE, destination)
    return destination


def test_f_doc_001_extracts_required_sections_and_stable_ids(tmp_path: Path) -> None:
    resume = read_resume(_copy_fixture(tmp_path))

    assert resume.section_names == ("Experience", "Projects", "Education", "Technical Skills")
    assert all(key.startswith("p-") for key in resume.paragraphs_by_id)
    assert len(resume.paragraphs_by_id) == len(set(resume.paragraphs_by_id))


def test_f_rwr_001_contact_block_is_absent_from_model_payload(tmp_path: Path) -> None:
    resume = read_resume(_copy_fixture(tmp_path))
    serialized = json.dumps(resume.model_payload())

    assert "Jordan Example" not in serialized
    assert "jordan@example.invalid" not in serialized
    assert "202-555-0147" not in serialized
    assert "portfolio.example.invalid" not in serialized


@pytest.mark.parametrize(
    "value",
    ["person@example.invalid", "202-555-0199", "https://example.invalid/person"],
)
def test_contact_patterns_are_detected_and_redacted(value: str) -> None:
    assert contains_pii(value)
    assert value not in redact_pii(f"before {value} after")


def test_body_contact_data_is_redacted_and_nonrewritable(tmp_path: Path) -> None:
    path = tmp_path / "body-pii.docx"
    document = Document()
    document.add_paragraph("Fictional Person")
    document.add_paragraph("person@example.invalid")
    document.add_paragraph("Experience", style="Heading 1")
    document.add_paragraph("Contact 202-555-0199 for fictional details")
    document.add_paragraph("Education", style="Heading 1")
    document.add_paragraph("Example University")
    document.save(str(path))

    resume = read_resume(path)
    paragraph = next(iter(resume.paragraphs_by_id.values()))

    assert paragraph.rewritable is False
    assert "202-555-0199" not in paragraph.model_text
    assert "[REDACTED_PHONE]" in paragraph.model_text


def test_hyperlink_paragraph_is_preserved_but_not_rewritable(tmp_path: Path) -> None:
    resume = read_resume(_copy_fixture(tmp_path))
    hyperlink = next(
        paragraph
        for paragraph in resume.paragraphs_by_id.values()
        if "Fictional project reference" in paragraph.source_text
    )

    assert hyperlink.rewritable is False


@pytest.mark.parametrize("kind", ["no_contact", "no_experience", "no_education"])
def test_missing_required_structure_is_rejected(tmp_path: Path, kind: str) -> None:
    path = tmp_path / f"{kind}.docx"
    document = Document()
    if kind != "no_contact":
        document.add_paragraph("Fictional Person")
    if kind != "no_experience":
        document.add_paragraph("Experience", style="Heading 1")
        document.add_paragraph("Fictional role")
    if kind != "no_education":
        document.add_paragraph("Education", style="Heading 1")
        document.add_paragraph("Fictional school")
    document.save(str(path))

    with pytest.raises(ResumeStructureError):
        read_resume(path)


def test_unreadable_or_missing_docx_is_rejected(tmp_path: Path) -> None:
    with pytest.raises(ResumeStructureError, match="not found"):
        read_resume(tmp_path / "missing.docx")
    invalid = tmp_path / "invalid.docx"
    invalid.write_bytes(b"not a DOCX")
    with pytest.raises(ResumeStructureError, match="readable"):
        read_resume(invalid)
