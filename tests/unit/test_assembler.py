"""F-DOC template preservation, privacy scrub, shadow, and publication tests."""

from __future__ import annotations

import hashlib
import shutil
import zipfile
from datetime import UTC, datetime
from pathlib import Path

import pytest
from docx import Document

from auto_interner.documents import assembler as assembler_module
from auto_interner.documents.assembler import (
    DocumentAssemblyError,
    assemble_resume,
)
from auto_interner.documents.template_reader import ResumeDocument, read_resume
from auto_interner.models import Listing
from auto_interner.paths import OutputCollisionError, OutputPathPlanner
from auto_interner.rewriting.service import validate_rewrite

pytestmark = [pytest.mark.unit, pytest.mark.privacy]
NOW = datetime(2027, 1, 2, 3, 4, tzinfo=UTC)
FIXTURE = (
    Path(__file__).resolve().parents[2]
    / "src"
    / "auto_interner"
    / "demo_data"
    / "fictional_base_resume.docx"
)


def _copy_base(tmp_path: Path) -> Path:
    base = tmp_path / "baseplate" / "base_resume.docx"
    base.parent.mkdir()
    shutil.copyfile(FIXTURE, base)
    return base


def _paragraph_id(resume: ResumeDocument, phrase: str) -> str:
    return next(
        paragraph.paragraph_id
        for paragraph in resume.paragraphs_by_id.values()
        if phrase in paragraph.source_text
    )


def _plan(resume: ResumeDocument, *, reorder: bool = True) -> object:
    order = (
        ["Technical Skills", "Experience", "Projects", "Education"]
        if reorder
        else list(resume.section_names)
    )
    return validate_rewrite(
        resume,
        {
            "section_order": order,
            "replacements": [
                {
                    "paragraph_id": _paragraph_id(resume, "reduced fictional review time"),
                    "replacement": (
                        "Using Python validation tools, reduced fictional review time by 30% "
                        "across 12,000 records."
                    ),
                }
            ],
        },
    )


def _destination(tmp_path: Path) -> Path:
    listing = Listing(
        id="fictional-listing",
        company_name="Fictional Systems",
        title="Software Engineering Intern",
        url="https://example.invalid/job",
        locations=("Denver, CO",),
        active=True,
    )
    planner = OutputPathPlanner(tmp_path / "data", 2027)
    path_plan = planner.plan(listing, generated_at=NOW)
    return planner.prepare(path_plan)


def test_f_doc_002_003_004_patches_copy_reorders_and_keeps_source_immutable(
    tmp_path: Path,
) -> None:
    base = _copy_base(tmp_path)
    source_hash = hashlib.sha256(base.read_bytes()).hexdigest()
    resume = read_resume(base)
    destination = _destination(tmp_path)

    result = assemble_resume(resume, _plan(resume), destination)

    assert result.written
    assert result.output_path == destination
    assert hashlib.sha256(base.read_bytes()).hexdigest() == source_hash
    generated = read_resume(destination)
    assert generated.section_names == (
        "Technical Skills",
        "Experience",
        "Projects",
        "Education",
    )
    assert "Using Python validation tools" in "\n".join(
        paragraph.source_text for paragraph in generated.paragraphs_by_id.values()
    )


def test_f_doc_005_page_geometry_styles_and_bold_runs_are_preserved(tmp_path: Path) -> None:
    base = _copy_base(tmp_path)
    resume = read_resume(base)
    destination = _destination(tmp_path)
    assemble_resume(resume, _plan(resume, reorder=False), destination)

    source = Document(str(base))
    generated = Document(str(destination))
    source_section = source.sections[0]
    generated_section = generated.sections[0]
    assert (
        generated_section.page_width,
        generated_section.page_height,
        generated_section.top_margin,
        generated_section.left_margin,
    ) == (
        source_section.page_width,
        source_section.page_height,
        source_section.top_margin,
        source_section.left_margin,
    )
    source_role = next(p for p in source.paragraphs if "Northstar Fictional" in p.text)
    generated_role = next(p for p in generated.paragraphs if "Northstar Fictional" in p.text)
    assert generated_role.style.name == source_role.style.name
    assert generated_role.runs[0].bold == source_role.runs[0].bold is True


def test_f_doc_006_external_hyperlinks_are_preserved(tmp_path: Path) -> None:
    base = _copy_base(tmp_path)
    resume = read_resume(base)
    destination = _destination(tmp_path)
    assemble_resume(resume, _plan(resume), destination)

    with zipfile.ZipFile(destination) as package:
        relationships = package.read("word/_rels/document.xml.rels").decode()

    assert "https://portfolio.example.invalid" in relationships
    assert "https://example.invalid/project" in relationships


def test_f_doc_007_private_metadata_rsids_and_zip_timestamps_are_scrubbed(
    tmp_path: Path,
) -> None:
    base = _copy_base(tmp_path)
    resume = read_resume(base)
    destination = _destination(tmp_path)
    assemble_resume(resume, _plan(resume), destination)

    with zipfile.ZipFile(destination) as package:
        core = package.read("docProps/core.xml").decode()
        word_xml = b"".join(
            package.read(name)
            for name in package.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )
        timestamps = {item.date_time for item in package.infolist()}

    assert "Fictional Fixture Builder" not in core
    assert "Jordan Example" not in core
    assert b"rsid" not in word_xml
    assert timestamps == {(1980, 1, 1, 0, 0, 0)}


def test_f_doc_008_existing_destination_is_never_overwritten(tmp_path: Path) -> None:
    base = _copy_base(tmp_path)
    resume = read_resume(base)
    destination = _destination(tmp_path)
    destination.write_bytes(b"existing artifact")

    with pytest.raises(OutputCollisionError):
        assemble_resume(resume, _plan(resume), destination)

    assert destination.read_bytes() == b"existing artifact"


def test_f_doc_009_shadow_mode_writes_nothing_and_creates_no_directory(tmp_path: Path) -> None:
    base = _copy_base(tmp_path)
    resume = read_resume(base)
    destination = tmp_path / "missing" / "company" / "role_01-02-27.docx"

    result = assemble_resume(resume, _plan(resume), destination, shadow_mode=True)

    assert result.output_path == destination
    assert result.written is False
    assert not destination.parent.exists()


def test_f_doc_010_unprepared_or_symbolic_output_directory_is_rejected(
    tmp_path: Path,
) -> None:
    base = _copy_base(tmp_path)
    resume = read_resume(base)

    with pytest.raises(DocumentAssemblyError, match="prepared"):
        assemble_resume(resume, _plan(resume), tmp_path / "missing" / "resume.docx")


def test_f_doc_011_base_change_after_extraction_is_rejected(tmp_path: Path) -> None:
    base = _copy_base(tmp_path)
    resume = read_resume(base)
    base.write_bytes(base.read_bytes() + b"changed")

    with pytest.raises(DocumentAssemblyError, match="changed"):
        assemble_resume(resume, _plan(resume), _destination(tmp_path))


def test_f_doc_012_validation_failure_leaves_no_artifact_or_temp_files(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    base = _copy_base(tmp_path)
    resume = read_resume(base)
    destination = _destination(tmp_path)

    def fail_validation(*args: object, **kwargs: object) -> None:
        del args, kwargs
        raise DocumentAssemblyError("fictional validation failure")

    monkeypatch.setattr(assembler_module, "_validate_output", fail_validation)

    with pytest.raises(DocumentAssemblyError, match="fictional"):
        assemble_resume(resume, _plan(resume), destination)

    assert not destination.exists()
    assert list(destination.parent.glob(".auto-interner-*.docx")) == []


def test_f_doc_013_generated_package_reopens_and_has_expected_filename(tmp_path: Path) -> None:
    base = _copy_base(tmp_path)
    resume = read_resume(base)
    destination = _destination(tmp_path)

    assemble_resume(resume, _plan(resume), destination)

    assert destination.name == "engineering-software_01-02-27.docx"
    reopened = Document(str(destination))
    assert any("Jordan Example" in paragraph.text for paragraph in reopened.paragraphs)
