"""Render a Selection into the user's own resume template.

The corpus engine decides *what* appears. This decides *how it looks*, and the
answer is: exactly like the resume the user already has.

Building a DOCX from scratch with python-docx produces a document that is
correct and looks nothing like theirs — different font, different margins,
different heading rules, no tab stops. A tailored resume that does not match the
one they have been sending is a new document they now have to proofread.

So this copies the base file and deletes the paragraphs selection left out.
Everything that survives keeps its original run properties, styles, tab stops,
and hyperlinks, because none of it was touched.

`documents/assembler.py` reorders sections and replaces paragraph text for the
rewrite pipeline. It has no deletion path, which is why this is separate rather
than an extra flag there.
"""

from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.text.paragraph import Paragraph

from auto_interner.corpus.formatting import (
    BASE_PAGE_WIDTH_EMU,
    BASE_SIDE_MARGIN_EMU,
    emu,
)
from auto_interner.corpus.selection import Selection
from auto_interner.documents.template_reader import ResumeDocument, heading_name


class TemplateAssemblyError(RuntimeError):
    """The base template could not be reduced to the selection."""


@dataclass(frozen=True, slots=True)
class TemplateResult:
    destination: Path
    kept: int
    removed: int
    sections_kept: tuple[str, ...]

    def describe(self) -> str:
        return (
            f"{self.kept} paragraphs kept, {self.removed} removed, "
            f"sections: {', '.join(self.sections_kept)}"
        )


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W = f"{{{W}}}"

# A repository link belongs on its project's title line, pushed to the right
# margin -- the position experience and education already use for location and
# dates. Anything else spends a line of a one-page budget on a URL.
#
# The base resume reached that position with 84 literal spaces, which is not a
# position at all: it holds only at one font size and one margin, and here it
# overflowed and wrapped the link onto its own line anyway. A right tab stop is
# the same intent expressed so it cannot drift.
_LINK_PADDING = re.compile(r"^\s{4,}$")


def _drop(paragraph: Paragraph) -> None:
    """Remove a paragraph from its parent, leaving surrounding XML intact."""
    element = paragraph._p
    element.getparent().remove(element)


def _right_tab_at(paragraph: Paragraph, position: int) -> None:
    """Guarantee a right-aligned tab stop at the right margin."""
    properties = paragraph._p.get_or_add_pPr()
    tabs = properties.find(f"{_W}tabs")
    if tabs is None:
        tabs = properties.makeelement(f"{_W}tabs", {})
        properties.append(tabs)
    for stop in tabs:
        if stop.get(f"{_W}val") == "right" and stop.get(f"{_W}pos") == str(position):
            return
    stop = tabs.makeelement(f"{_W}tab", {f"{_W}val": "right", f"{_W}pos": str(position)})
    tabs.append(stop)


def inline_trailing_link(paragraph: Paragraph, right_margin_twips: int) -> bool:
    """Move a padded trailing hyperlink onto the title line, right-aligned.

    Applies only to a paragraph holding exactly one hyperlink that is separated
    from the title by whitespace alone. That excludes the contact block, where
    several links are punctuation-separated inside a real sentence, and any
    prose that merely happens to end in a link.

    Whitespace is the only thing added or removed. No claim, character or run
    property changes, which `test_link_inlining_changes_no_words` asserts.
    """
    links = paragraph._p.findall(f"{_W}hyperlink")
    if len(links) != 1:
        return False
    link = links[0]
    children = list(paragraph._p)
    index = children.index(link)

    # Everything between the title and the link that exists only to position it:
    # literal spaces, tab characters, and the text-less runs Word leaves behind.
    # A tab already being present does not mean the line is correct -- the base
    # resume had a tab *and* 84 spaces *and* no right stop, so the tab fell on a
    # default stop and the spaces pushed the link onto the next line anyway.
    spacing: list[BaseOxmlElement] = []
    for node in reversed(children[:index]):
        if node.tag != f"{_W}r":
            break
        text = node.xpath("string(.)")
        if text.strip():
            break
        spacing.append(node)
    if not spacing:
        return False

    literal = "".join(n.xpath("string(.)") for n in spacing)
    tabs = sum(len(n.findall(f"{_W}tab")) for n in spacing)
    already_right = any(
        stop.get(f"{_W}val") == "right"
        for stop in paragraph._p.findall(f"{_W}pPr/{_W}tabs/{_W}tab")
    )
    padded = bool(_LINK_PADDING.match(literal))
    if not padded and tabs == 1 and already_right and len(spacing) == 1:
        return False  # already correct; keep the rule idempotent
    if not padded and tabs == 0:
        return False  # a single space before a link is prose, not a layout

    keeper = spacing[0]
    for node in spacing[1:]:
        paragraph._p.remove(node)
    for child in list(keeper):
        if child.tag != f"{_W}rPr":
            keeper.remove(child)
    keeper.append(keeper.makeelement(f"{_W}tab", {}))

    for node in children[index + 1 :]:
        if node.tag == f"{_W}r" and not node.xpath("string(.)").strip():
            paragraph._p.remove(node)

    _right_tab_at(paragraph, right_margin_twips)
    return True


def inline_trailing_links(document: DocxDocument) -> int:
    """Apply the title-line link rule to every paragraph that qualifies."""
    section = document.sections[0]
    width = (
        emu(section.page_width, BASE_PAGE_WIDTH_EMU)
        - emu(section.left_margin, BASE_SIDE_MARGIN_EMU)
        - emu(section.right_margin, BASE_SIDE_MARGIN_EMU)
    )
    right_margin_twips = round(width / 914400 * 1440)
    return sum(
        inline_trailing_link(paragraph, right_margin_twips) for paragraph in document.paragraphs
    )


def _selected_texts(selection: Selection) -> set[str]:
    """Exact source text of every bullet that survived selection.

    Matching on text rather than paragraph index because the corpus was built
    from this same document, so the strings are identical by construction, and
    an index would break the moment anything upstream reorders.
    """
    return {b.bullet.text.strip() for sb in selection.blocks for b in sb.bullets}


def _selected_headers(selection: Selection, source: ResumeDocument) -> set[str]:
    """Block titles and org lines that must survive, matched to source text."""
    wanted: set[str] = set()
    for sb in selection.blocks:
        for field in (sb.block.title, sb.block.org):
            if field:
                wanted.add(field.strip())
    # Resume headers carry their metadata after a separator, and the separator
    # differs by section: `Org<TAB>Location` in experience, `Project | Stack` in
    # projects. Splitting on tab alone silently deleted every project title,
    # leaving its bullets orphaned under the section heading.
    resolved: set[str] = set()
    for section in source.sections:
        for p in section.paragraphs:
            left = re.split(r"[\t|]", p.source_text, maxsplit=1)[0].strip()
            if left in wanted:
                resolved.add(p.source_text.strip())
    return resolved


def assemble_from_template(
    source_resume: ResumeDocument,
    selection: Selection,
    destination: Path,
    *,
    keep_empty_sections: bool = False,
) -> TemplateResult:
    """Copy the base resume and delete everything selection did not keep.

    The contact block and section headings are always preserved — they carry the
    document's identity and its visual structure, and neither is a corpus block.
    """
    if destination.exists():
        raise TemplateAssemblyError(f"refusing to overwrite {destination}")
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source_resume.source_path, destination)

    document = Document(str(destination))
    keep_text = _selected_texts(selection) | _selected_headers(selection, source_resume)

    kept = removed = 0
    seen_heading = False
    survivors_by_section: dict[str, int] = {}
    current: str | None = None
    doomed: list[Paragraph] = []

    for paragraph in document.paragraphs:
        heading = heading_name(paragraph)
        if heading is not None:
            current = heading
            survivors_by_section.setdefault(current, 0)
            seen_heading = True
            kept += 1
            continue
        text = paragraph.text.strip()
        if not seen_heading:
            kept += 1  # contact block, above the first heading
            continue
        if not text:
            kept += 1  # spacing
            continue
        if text in keep_text:
            kept += 1
            if current:
                survivors_by_section[current] += 1
        else:
            doomed.append(paragraph)
            removed += 1

    for paragraph in doomed:
        _drop(paragraph)

    # A bullet with no text renders as a floating dot. Style name is not enough
    # to find them: Word bullets are usually numbering properties (`w:numPr`) on
    # an otherwise normal paragraph, so a name check misses every one.
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip():
            continue
        numbered = (
            paragraph._p.find(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
            )
            is not None
        )
        style_name = paragraph.style.name if paragraph.style is not None else None
        if numbered or "List" in (style_name or ""):
            _drop(paragraph)
            removed += 1

    if not keep_empty_sections:
        for paragraph in list(document.paragraphs):
            name = heading_name(paragraph)
            if name is not None and survivors_by_section.get(name, 0) == 0:
                _drop(paragraph)

    inline_trailing_links(document)

    if kept == 0:
        raise TemplateAssemblyError("selection kept nothing; refusing to write an empty resume")

    document.save(str(destination))
    return TemplateResult(
        destination=destination,
        kept=kept,
        removed=removed,
        sections_kept=tuple(n for n, c in survivors_by_section.items() if c),
    )
