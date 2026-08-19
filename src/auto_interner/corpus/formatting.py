"""Fit a finished document onto one page by typography alone.

Selection decides *what* appears and answers to a line budget. That budget is an
estimate, and an estimate that is slightly wrong spills a résumé onto a second
page — where the spilled part is usually a single orphaned line, which reads
worse than either a full page or a real two-page résumé.

This closes that gap from the other end: given a document that nearly fits,
tighten margins, spacing and type until it does. Content is never touched. If
the floors are reached and it still does not fit, that is reported rather than
papered over, because the honest remedy at that point is to cut a block.

Order matters. Adjustments run cheapest-looking first — dead space, then
paragraph spacing, then margins, then line spacing, and type size last. A reader
notices 9.5pt type long before they notice a 0.6in margin.

No new dependency: the capacity model is the same characters-per-line
arithmetic `Bullet.cost` already uses, so the two agree by construction.
"""

from __future__ import annotations

from dataclasses import dataclass, replace
from math import ceil, floor

from docx.document import Document as DocxDocument
from docx.shared import Inches, Pt

_EMU_PER_INCH = 914400
_POINTS_PER_INCH = 72

# A serif face at a given size averages close to half its point size in width.
# 7in usable at 11pt gives ~92 characters, which is the 95 that `Bullet.cost`
# assumes -- the two models have to agree or selection and fitting fight.
_AVERAGE_CHAR_WIDTH_EM = 0.5

# Word's single spacing is about 1.15x the type size before any explicit rule.
_SINGLE_SPACING = 1.15


@dataclass(frozen=True, slots=True)
class Typography:
    """Everything this module is allowed to change."""

    font_pt: float
    margin_in: float
    line_spacing: float
    space_after_pt: float

    def describe(self) -> str:
        return (
            f"{self.font_pt:g}pt / {self.margin_in:g}in margins / "
            f"{self.line_spacing:g} spacing / {self.space_after_pt:g}pt after"
        )


@dataclass(frozen=True, slots=True)
class Floors:
    """How far this is permitted to go.

    An unreadable résumé that fits is not a success. 9.5pt and 0.5in are about
    where a recruiter starts to notice; below that the document announces that
    it was squeezed, which is worse than running to a second page honestly.
    """

    font_pt: float = 9.5
    margin_in: float = 0.5
    line_spacing: float = 0.95
    space_after_pt: float = 0.0


@dataclass(frozen=True, slots=True)
class FitResult:
    fitted: bool
    steps: tuple[str, ...]
    start: Typography
    final: Typography
    lines: int
    capacity: int

    def describe(self) -> str:
        head = "fits one page" if self.fitted else "STILL OVER after every allowed step"
        body = "\n".join(f"  {s}" for s in self.steps) or "  no change needed"
        return (
            f"{head}: {self.lines} lines into {self.capacity}\n"
            f"{body}\n"
            f"  {self.start.describe()}\n"
            f"  -> {self.final.describe()}"
        )


def current_typography(document: DocxDocument) -> Typography:
    """Read the document's present settings."""
    section = document.sections[0]
    style = document.styles["Normal"]
    size = style.font.size
    spacing = style.paragraph_format.line_spacing
    after = style.paragraph_format.space_after
    return Typography(
        font_pt=size.pt if size is not None else 11.0,
        margin_in=min(section.left_margin, section.right_margin) / _EMU_PER_INCH,
        line_spacing=float(spacing) if isinstance(spacing, (int, float)) else 1.0,
        space_after_pt=after.pt if after is not None else 0.0,
    )


# Estimated capacity is deliberately under the arithmetic maximum.
#
# The model was calibrated against LibreOffice, which agreed with it exactly --
# and the documents still spilled a single bullet onto a second page in Word.
# The two lay out the same file differently: font metrics, hyphenation and
# widow/orphan control are not the same implementation, and the arithmetic here
# cannot be right for both.
#
# The two errors are not symmetric. Under-filling costs a little white space.
# Overflowing costs a whole extra sheet carrying one orphaned line, which is the
# single worst thing a résumé can look like. So the estimate is biased toward
# the recoverable failure, and the renderer that matters is the one the reader
# opens it in.
_SAFETY = 0.88


def capacity(
    document: DocxDocument, typography: Typography, *, safety: float = _SAFETY
) -> int:
    """How many rendered lines one page holds at these settings."""
    section = document.sections[0]
    usable_pt = (
        (section.page_height - section.top_margin - section.bottom_margin)
        / _EMU_PER_INCH
        * _POINTS_PER_INCH
    )
    per_line = typography.font_pt * _SINGLE_SPACING * typography.line_spacing
    per_line += typography.space_after_pt
    return max(1, floor(usable_pt * safety / per_line))


def estimate_lines(document: DocxDocument, typography: Typography) -> int:
    """Rendered line count, wrapping long paragraphs the way the page will."""
    section = document.sections[0]
    usable_in = (
        section.page_width - section.left_margin - section.right_margin
    ) / _EMU_PER_INCH
    chars = max(
        20,
        int(usable_in * _POINTS_PER_INCH / (typography.font_pt * _AVERAGE_CHAR_WIDTH_EM)),
    )
    total = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        total += 1 if not text else ceil(len(text) / chars)
    return total


def apply(document: DocxDocument, typography: Typography) -> None:
    """Write the settings onto the document.

    Sizes are set on the `Normal` style and on every run that carries an
    explicit size, because a run-level size overrides the style and a résumé
    template is full of them.
    """
    for section in document.sections:
        section.left_margin = Inches(typography.margin_in)
        section.right_margin = Inches(typography.margin_in)
        section.top_margin = Inches(typography.margin_in)
        section.bottom_margin = Inches(typography.margin_in)

    normal = document.styles["Normal"]
    base = normal.font.size.pt if normal.font.size is not None else 11.0
    scale = typography.font_pt / base if base else 1.0
    normal.font.size = Pt(typography.font_pt)
    normal.paragraph_format.line_spacing = typography.line_spacing
    normal.paragraph_format.space_after = Pt(typography.space_after_pt)

    for paragraph in document.paragraphs:
        paragraph.paragraph_format.space_after = Pt(typography.space_after_pt)
        paragraph.paragraph_format.line_spacing = typography.line_spacing
        for run in paragraph.runs:
            if run.font.size is not None:
                # Scale rather than flatten: the name and section headings are
                # deliberately larger, and setting every run to one size erases
                # the document's hierarchy along with its overflow.
                run.font.size = Pt(max(typography.font_pt, run.font.size.pt * scale))


def drop_empty_paragraphs(document: DocxDocument, *, keep: int = 0) -> int:
    """Remove blank paragraphs, the cheapest space there is.

    Consecutive blanks are pure dead space. `keep` blanks are preserved between
    populated paragraphs so sections do not collide.
    """
    removed = 0
    run_length = 0
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip():
            run_length = 0
            continue
        run_length += 1
        if run_length > keep:
            paragraph._p.getparent().remove(paragraph._p)
            removed += 1
    return removed


def fit_to_one_page(
    document: DocxDocument,
    *,
    floors: Floors | None = None,
    font_step: float = 0.5,
    margin_step: float = 0.1,
) -> FitResult:
    """Tighten typography until the document fits one page, or floors are hit."""
    floors = floors or Floors()
    start = current_typography(document)
    typography = start
    steps: list[str] = []

    removed = drop_empty_paragraphs(document, keep=0)
    if removed:
        steps.append(f"removed {removed} blank paragraph(s)")

    def over() -> bool:
        return estimate_lines(document, typography) > capacity(document, typography)

    # Cheapest-looking first. Paragraph spacing and margins are invisible to a
    # reader in a way that type size is not.
    if over() and typography.space_after_pt > floors.space_after_pt:
        typography = replace(typography, space_after_pt=floors.space_after_pt)
        steps.append(f"space after -> {floors.space_after_pt:g}pt")

    while over() and typography.margin_in - margin_step >= floors.margin_in:
        typography = replace(typography, margin_in=round(typography.margin_in - margin_step, 2))
        steps.append(f"margins -> {typography.margin_in:g}in")

    while over() and round(typography.line_spacing - 0.05, 2) >= floors.line_spacing:
        typography = replace(typography, line_spacing=round(typography.line_spacing - 0.05, 2))
        steps.append(f"line spacing -> {typography.line_spacing:g}")

    while over() and typography.font_pt - font_step >= floors.font_pt:
        typography = replace(typography, font_pt=round(typography.font_pt - font_step, 2))
        steps.append(f"font -> {typography.font_pt:g}pt")

    apply(document, typography)
    lines = estimate_lines(document, typography)
    room = capacity(document, typography)
    return FitResult(
        fitted=lines <= room,
        steps=tuple(steps),
        start=start,
        final=typography,
        lines=lines,
        capacity=room,
    )
