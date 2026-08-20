"""Read resume structure while separating contact data from model-safe content."""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

_KNOWN_HEADINGS = frozenset(
    {
        "experience",
        "professional experience",
        "work experience",
        "education",
        "skills",
        "technical skills",
        "projects",
        "certifications",
        "awards",
        "leadership",
        "activities",
        "summary",
    }
)
_REQUIRED_HEADING_GROUPS = (
    frozenset({"experience", "professional experience", "work experience"}),
    frozenset({"education"}),
)

# Only these sections are offered to the rewriter. The contact block is already
# outside `sections` entirely and never reaches a model, so the tailorable
# surface is experience and education alone.
#
# Everything else -- projects, skills, coursework -- keeps its heading so
# `section_order` can still place it, but its paragraph text is neither sent nor
# replaceable. That is most of the payload for no tailoring value: a skills line
# cannot be rephrased without either inventing a technology or saying the same
# thing differently, and both are rejected downstream anyway.
_REWRITABLE_SECTION_HEADINGS = frozenset(
    {"experience", "professional experience", "work experience", "education"}
)
_EMAIL_PATTERN = re.compile(r"(?<![\w.+-])[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}(?![\w-])")
_PHONE_PATTERN = re.compile(r"(?<!\w)(?:\+?1[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)\d{3}[\s.-]?\d{4}(?!\w)")
_URL_PATTERN = re.compile(r"(?i)\b(?:https?://|www\.)[^\s<>()]+")


class ResumeStructureError(ValueError):
    """The base DOCX lacks the stable structure required for safe rewriting."""


@dataclass(frozen=True, slots=True)
class ResumeParagraph:
    """One stable paragraph locator and its model-safe representation."""

    paragraph_id: str
    source_text: str
    model_text: str
    rewritable: bool


@dataclass(frozen=True, slots=True)
class ResumeSection:
    """A detected named resume section in document order."""

    name: str
    paragraphs: tuple[ResumeParagraph, ...]


@dataclass(frozen=True, slots=True)
class ResumeDocument:
    """Immutable extraction result; contact data never enters `model_payload`."""

    source_path: Path
    source_sha256: str
    contact_text: tuple[str, ...]
    sections: tuple[ResumeSection, ...]

    @property
    def section_names(self) -> tuple[str, ...]:
        """Return section names in base-document order."""
        return tuple(section.name for section in self.sections)

    @property
    def paragraphs_by_id(self) -> dict[str, ResumeParagraph]:
        """Return stable paragraph lookup used by validation and assembly."""
        return {
            paragraph.paragraph_id: paragraph
            for section in self.sections
            for paragraph in section.paragraphs
        }

    def model_payload(self) -> dict[str, object]:
        """Return only sanitized, non-contact resume sections for a rewrite call.

        A section outside `_REWRITABLE_SECTION_HEADINGS` is named but empty. The
        name has to survive, because `validate_rewrite` requires `section_order`
        to list every section exactly once; the paragraph text does not, because
        nothing in that section may be replaced. Dropping it is the single
        largest reduction available to this request.
        """
        sections: list[dict[str, object]] = []
        for section in self.sections:
            in_scope = section.name.strip().casefold() in _REWRITABLE_SECTION_HEADINGS
            sections.append(
                {
                    "name": section.name,
                    "rewritable_section": in_scope,
                    "paragraphs": [
                        {
                            "paragraph_id": paragraph.paragraph_id,
                            "text": paragraph.model_text,
                            "rewritable": paragraph.rewritable,
                        }
                        for paragraph in section.paragraphs
                    ]
                    if in_scope
                    else [],
                }
            )
        return {"sections": sections}


def contains_pii(text: str) -> bool:
    """Return whether text contains a contact-data pattern."""
    return any(pattern.search(text) for pattern in (_EMAIL_PATTERN, _PHONE_PATTERN, _URL_PATTERN))


def redact_pii(text: str) -> str:
    """Replace contact patterns with stable non-sensitive tokens."""
    sanitized = _EMAIL_PATTERN.sub("[REDACTED_EMAIL]", text)
    sanitized = _PHONE_PATTERN.sub("[REDACTED_PHONE]", sanitized)
    return _URL_PATTERN.sub("[REDACTED_URL]", sanitized)


def _contains_hyperlink(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:hyperlink"))


def heading_name(paragraph: Paragraph) -> str | None:
    """Return a recognized section heading without altering its display text."""
    text = " ".join(paragraph.text.split())
    if not text:
        return None
    normalized = text.casefold().rstrip(":")
    style_name = paragraph.style.name.casefold() if paragraph.style is not None else ""
    if style_name.startswith("heading") or normalized in _KNOWN_HEADINGS:
        return text.rstrip(":")
    return None


def _validate_required_sections(sections: list[ResumeSection]) -> None:
    normalized = {section.name.casefold().rstrip(":") for section in sections}
    for choices in _REQUIRED_HEADING_GROUPS:
        if normalized.isdisjoint(choices):
            expected = " or ".join(sorted(choices))
            raise ResumeStructureError(f"Base resume is missing required section: {expected}")


def read_resume(path: Path) -> ResumeDocument:
    """Extract a DOCX with a hard contact/model boundary and stable paragraph IDs."""
    if not path.is_file():
        raise ResumeStructureError("Base resume DOCX was not found")
    source_bytes = path.read_bytes()
    source_hash = hashlib.sha256(source_bytes).hexdigest()
    try:
        document: DocxDocument = Document(str(path))
    except Exception as exc:
        raise ResumeStructureError("Base resume is not a readable DOCX") from exc

    contact: list[str] = []
    sections: list[ResumeSection] = []
    current_name: str | None = None
    current_paragraphs: list[ResumeParagraph] = []
    for index, paragraph in enumerate(document.paragraphs):
        heading = heading_name(paragraph)
        if heading is not None:
            if current_name is not None:
                sections.append(ResumeSection(current_name, tuple(current_paragraphs)))
            current_name = heading
            current_paragraphs = []
            continue
        text = paragraph.text.strip()
        if current_name is None:
            if text:
                contact.append(text)
            continue
        if not text:
            continue
        has_pii = contains_pii(text)
        has_hyperlink = _contains_hyperlink(paragraph)
        in_scope = current_name.strip().casefold() in _REWRITABLE_SECTION_HEADINGS
        current_paragraphs.append(
            ResumeParagraph(
                paragraph_id=f"p-{index}",
                source_text=text,
                model_text=redact_pii(text),
                rewritable=in_scope and not has_pii and not has_hyperlink,
            )
        )
    if current_name is not None:
        sections.append(ResumeSection(current_name, tuple(current_paragraphs)))
    if not contact:
        raise ResumeStructureError("Base resume must have a contact block before its first section")
    if not sections:
        raise ResumeStructureError("Base resume has no recognized sections")
    _validate_required_sections(sections)
    if len({section.name.casefold() for section in sections}) != len(sections):
        raise ResumeStructureError("Base resume section headings must be unique")
    return ResumeDocument(path, source_hash, tuple(contact), tuple(sections))
