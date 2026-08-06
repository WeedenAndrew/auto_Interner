"""Template-preserving DOCX patching and collision-safe publication."""

from __future__ import annotations

import hashlib
import os
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import cast

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from lxml import etree  # type: ignore[import-untyped]

from auto_interner.documents.template_reader import (
    ResumeDocument,
    ResumeStructureError,
    heading_name,
    read_resume,
)
from auto_interner.paths import OutputCollisionError
from auto_interner.rewriting.service import ValidatedRewritePlan

_CORE_PROPERTIES = "docProps/core.xml"
_CUSTOM_PROPERTIES = "docProps/custom.xml"
_PACKAGE_RELATIONSHIPS = "_rels/.rels"
_CONTENT_TYPES = "[Content_Types].xml"
_FIXED_W3CDTF = "2000-01-01T00:00:00Z"


class DocumentAssemblyError(RuntimeError):
    """Generated DOCX failed preservation, validation, or atomic publication."""


@dataclass(frozen=True, slots=True)
class AssemblyResult:
    """Actual or shadow-mode result for one validated output plan."""

    output_path: Path
    written: bool


def _source_hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _geometry(document: DocxDocument) -> tuple[tuple[int | None, ...], ...]:
    return tuple(
        (
            section.page_width,
            section.page_height,
            section.top_margin,
            section.right_margin,
            section.bottom_margin,
            section.left_margin,
            section.header_distance,
            section.footer_distance,
        )
        for section in document.sections
    )


def _replace_paragraph_text(paragraph: Paragraph, replacement: str) -> None:
    if paragraph._p.xpath(".//w:hyperlink"):
        raise DocumentAssemblyError("Validated plan targeted a hyperlink paragraph")
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(replacement)
        return
    runs[0].text = replacement
    for run in runs[1:]:
        run.text = ""


def _reorder_sections(document: DocxDocument, section_order: tuple[str, ...]) -> None:
    body = document.element.body
    children = [child for child in body.iterchildren() if child.tag.rsplit("}", 1)[-1] != "sectPr"]
    heading_by_node: dict[object, str] = {}
    for paragraph in document.paragraphs:
        name = heading_name(paragraph)
        if name is not None:
            heading_by_node[paragraph._p] = name
    starts = [
        (index, heading_by_node[node])
        for index, node in enumerate(children)
        if node in heading_by_node
    ]
    if tuple(name for _, name in starts) == section_order:
        return
    blocks: dict[str, list[object]] = {}
    for position, (start, name) in enumerate(starts):
        end = starts[position + 1][0] if position + 1 < len(starts) else len(children)
        blocks[name] = children[start:end]
    if set(blocks) != set(section_order):
        raise DocumentAssemblyError("DOCX section blocks do not match the validated plan")
    for nodes in blocks.values():
        for node in nodes:
            body.remove(node)
    section_properties = body.sectPr
    insert_at = body.index(section_properties) if section_properties is not None else len(body)
    for name in section_order:
        for node in blocks[name]:
            body.insert(insert_at, node)
            insert_at += 1


def _scrub_xml(name: str, data: bytes) -> bytes:
    try:
        parser = etree.XMLParser(resolve_entities=False, no_network=True)
        root = etree.fromstring(data, parser=parser)
    except etree.XMLSyntaxError as exc:
        raise DocumentAssemblyError("DOCX contains malformed XML") from exc
    if name == _CORE_PROPERTIES:
        for element in root:
            local_name = element.tag.rsplit("}", 1)[-1]
            if local_name in {"created", "modified"}:
                element.text = _FIXED_W3CDTF
            elif local_name == "revision":
                element.text = "1"
            else:
                element.text = ""
    elif name == _PACKAGE_RELATIONSHIPS:
        for element in list(root):
            if str(element.attrib.get("Type", "")).endswith("/custom-properties"):
                root.remove(element)
    elif name == _CONTENT_TYPES:
        for element in list(root):
            if element.attrib.get("PartName") == "/docProps/custom.xml":
                root.remove(element)
    if name.startswith("word/") and name.endswith(".xml"):
        for element in root.iter():
            for attribute in list(element.attrib):
                if attribute.rsplit("}", 1)[-1].startswith("rsid"):
                    del element.attrib[attribute]
            for child in list(element):
                if child.tag.rsplit("}", 1)[-1].startswith("rsid"):
                    element.remove(child)
    return cast(bytes, etree.tostring(root, encoding="utf-8", xml_declaration=True))


def _scrub_package(source: Path, destination: Path) -> None:
    try:
        with (
            zipfile.ZipFile(source, "r") as incoming,
            zipfile.ZipFile(destination, "w", compression=zipfile.ZIP_DEFLATED) as outgoing,
        ):
            for item in incoming.infolist():
                if item.filename == _CUSTOM_PROPERTIES:
                    continue
                data = incoming.read(item)
                if item.filename.endswith(".xml") or item.filename.endswith(".rels"):
                    data = _scrub_xml(item.filename, data)
                clean = zipfile.ZipInfo(item.filename, date_time=(1980, 1, 1, 0, 0, 0))
                clean.compress_type = zipfile.ZIP_DEFLATED
                clean.external_attr = item.external_attr
                outgoing.writestr(clean, data)
    except (OSError, zipfile.BadZipFile) as exc:
        raise DocumentAssemblyError("DOCX package could not be scrubbed") from exc


def _validate_output(
    output: Path,
    source_document: DocxDocument,
    source_resume: ResumeDocument,
    plan: ValidatedRewritePlan,
) -> None:
    try:
        with zipfile.ZipFile(output) as package:
            if package.testzip() is not None:
                raise DocumentAssemblyError("Generated DOCX package failed CRC validation")
            names = set(package.namelist())
            if _CUSTOM_PROPERTIES in names:
                raise DocumentAssemblyError("Generated DOCX retained custom properties")
        generated = Document(str(output))
        extracted = read_resume(output)
    except (OSError, zipfile.BadZipFile, ResumeStructureError) as exc:
        raise DocumentAssemblyError("Generated DOCX could not be reopened and validated") from exc
    if _geometry(generated) != _geometry(source_document):
        raise DocumentAssemblyError("Generated DOCX changed page geometry")
    if extracted.section_names != plan.section_order:
        raise DocumentAssemblyError("Generated DOCX section order differs from the validated plan")
    generated_text = "\n".join(paragraph.text for paragraph in generated.paragraphs)
    for contact_line in source_resume.contact_text:
        if contact_line not in generated_text:
            raise DocumentAssemblyError("Generated DOCX lost contact-block content")
    for replacement in plan.replacements:
        if replacement.replacement not in generated_text:
            raise DocumentAssemblyError("Generated DOCX omitted a validated replacement")


def assemble_resume(
    source_resume: ResumeDocument,
    plan: ValidatedRewritePlan,
    destination: Path,
    *,
    shadow_mode: bool = False,
) -> AssemblyResult:
    """Patch a copied template, validate it, and atomically publish without overwrite."""
    if shadow_mode:
        return AssemblyResult(destination, written=False)
    if destination.exists() or destination.is_symlink():
        raise OutputCollisionError("generated resume already exists; refusing to overwrite it")
    if not destination.parent.is_dir() or destination.parent.is_symlink():
        raise DocumentAssemblyError("Output directory must be safely prepared before assembly")
    if _source_hash(source_resume.source_path) != source_resume.source_sha256:
        raise DocumentAssemblyError("Base resume changed after rewrite extraction")

    source_document = Document(str(source_resume.source_path))
    working_handle, working_name = tempfile.mkstemp(
        prefix=".auto-interner-working-", suffix=".docx", dir=destination.parent
    )
    scrubbed_handle, scrubbed_name = tempfile.mkstemp(
        prefix=".auto-interner-ready-", suffix=".docx", dir=destination.parent
    )
    os.close(working_handle)
    os.close(scrubbed_handle)
    working_path = Path(working_name)
    scrubbed_path = Path(scrubbed_name)
    try:
        paragraphs = source_document.paragraphs
        for replacement in plan.replacements:
            try:
                paragraph_index = int(replacement.paragraph_id.removeprefix("p-"))
                paragraph = paragraphs[paragraph_index]
            except (ValueError, IndexError) as exc:
                raise DocumentAssemblyError(
                    "Validated paragraph locator no longer resolves"
                ) from exc
            _replace_paragraph_text(paragraph, replacement.replacement)
        _reorder_sections(source_document, plan.section_order)
        source_document.save(str(working_path))
        _scrub_package(working_path, scrubbed_path)
        _validate_output(
            scrubbed_path,
            Document(str(source_resume.source_path)),
            source_resume,
            plan,
        )
        if _source_hash(source_resume.source_path) != source_resume.source_sha256:
            raise DocumentAssemblyError("Base resume changed during assembly")
        try:
            os.link(scrubbed_path, destination)
        except FileExistsError as exc:
            raise OutputCollisionError(
                "generated resume already exists; refusing to overwrite it"
            ) from exc
        return AssemblyResult(destination, written=True)
    finally:
        working_path.unlink(missing_ok=True)
        scrubbed_path.unlink(missing_ok=True)
